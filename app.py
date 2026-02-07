# app.py
# AI-based Legal Reference & Retrieval — full polished UI with theme selector and safe theme injection
# Question (user) -> RIGHT, Answer (assistant) -> LEFT
# Requires: rag.answer.answer_query (your RAG backend)

import os
from dotenv import load_dotenv
# Force load .env to override any stale/invalid system environment variables
load_dotenv(override=True)

import json
import time
import hashlib
import html
import markdown
from datetime import datetime
import streamlit as st

# import your backend answer function (must exist in project)
from rag.answer import answer_query

# -------------------------
# Page config
# -------------------------
st.set_page_config(page_title="AI Legal RAG — Polished UI", page_icon="⚖️", layout="wide", initial_sidebar_state="expanded")

# -------------------------
# Paths
# -------------------------
DATA_DIR = "data"
USERS_FP = os.path.join(DATA_DIR, "users.json")
CONVO_DIR = os.path.join(DATA_DIR, "conversations")
DOCS_DIR = os.path.join(DATA_DIR, "docs")
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(CONVO_DIR, exist_ok=True)
os.makedirs(DOCS_DIR, exist_ok=True)

# -------------------------
# Utilities
# -------------------------
def sha256(s: str) -> str:
    return hashlib.sha256(s.encode("utf-8")).hexdigest()

def read_json(fp, default):
    try:
        with open(fp, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default

def write_json(fp, data):
    with open(fp, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

def safe_truncate(text: str, n: int = 42) -> str:
    t = " ".join((text or "").split())
    return (t[:n] + "…") if len(t) > n else t

def now_ts():
    return int(time.time())

def fmt_time(ts):
    try:
        return datetime.fromtimestamp(ts).strftime("%b %d %H:%M")
    except Exception:
        return ""

def esc(s: str) -> str:
    return html.escape(s)

# -------------------------
# Users & Conversations
# -------------------------
def users_all(): return read_json(USERS_FP, [])
def users_save(u): write_json(USERS_FP, u)

def find_user(email: str):
    email = (email or "").strip().lower()
    for u in users_all():
        if u["email"].lower() == email:
            return u
    return None

def upsert_user(user):
    users = users_all()
    for i, u in enumerate(users):
        if u["email"].lower() == user["email"].lower():
            users[i] = user; users_save(users); return
    users.append(user); users_save(users)

def seed_demo_user():
    if not os.path.exists(USERS_FP):
        users_save([{"email":"demo@legal.ai","name":"Demo User","password_hash":sha256("demo1234")}])
seed_demo_user()

def convo_path(email: str) -> str:
    return os.path.join(CONVO_DIR, f"{email.lower().replace('@','_at_')}.json")

def convos_load(email: str):
    return read_json(convo_path(email), [])

def convos_save(email: str, convos):
    write_json(convo_path(email), convos)

def convos_ensure(email: str):
    if not os.path.exists(convo_path(email)):
        convos_save(email, [])
    return convos_load(email)

def convo_create(email: str, title: str = "New chat"):
    convos = convos_ensure(email)
    cid = f"c_{int(time.time()*1000)}"
    convo = {
        "id": cid,
        "title": title,
        "updatedAt": now_ts(),
        "messages": [
            {"role":"assistant","content":"Hi! Ask me anything about the law ⚖️","ts": time.time()}
        ],
    }
    convos.insert(0, convo)
    convos_save(email, convos)
    return cid

def convo_delete(email: str, cid: str):
    convs = convos_ensure(email)
    convs = [c for c in convs if c["id"] != cid]
    convos_save(email, convs)

def convo_append_msg(email: str, cid: str, role: str, content: str, suggestions: list = None):
    """
    Append a message to conversation `cid` for `email`.
    Ensures the conversation file exists, updates the convo, and saves it.
    """
    convs = convos_ensure(email)           # ensure 'convs' is defined
    for c in convs:
        if c["id"] == cid:
            msg = {"role": role, "content": content, "ts": time.time()}
            if suggestions:
                msg["suggestions"] = suggestions
            c["messages"].append(msg)
            c["updatedAt"] = now_ts()
            # if the convo title is still default/new, set it from user's first message
            if role == "user" and (c.get("title") in [None, "", "New chat"]) and len(c["messages"]) <= 3:
                c["title"] = safe_truncate(content, 52)
            break
    convos_save(email, convs)              # save the updated convos back to disk

def convo_clear(email: str, cid: str):
    convs = convos_ensure(email)
    for c in convs:
        if c["id"] == cid:
            c["messages"] = []
            c["updatedAt"] = now_ts()
            break
    convos_save(email, convs)

# -------------------------
# Auth helpers
# -------------------------
def login(email: str, password: str) -> bool:
    u = find_user(email)
    if not u or u["password_hash"] != sha256(password):
        return False
    st.session_state["user"] = {"email": u["email"], "name": u["name"]}
    convos_ensure(u["email"])
    if not st.session_state.get("active_cid"):
        convs = convos_load(u["email"])
        st.session_state["active_cid"] = convs[0]["id"] if convs else convo_create(u["email"])
    return True

def logout():
    st.session_state.pop("user", None)
    st.session_state.pop("active_cid", None)
    st.session_state.pop("route", None)
    st.rerun()

# -------------------------
# Themes + robust CSS generator
# -------------------------
# -------------------------
# Themes + robust CSS generator
# -------------------------
THEME_CONFIG = {
    "Midnight Purple": {
        "primary": "#6c5ce7",
        "primary_dark": "#4834d4",
        "bg_gradient": "linear-gradient(135deg, #0f0c29 0%, #302b63 50%, #24243e 100%)",
        "secondary_bg": "rgba(255, 255, 255, 0.05)",
        "text": "#e0e0e0"
    },
    "Deep Ocean Blue": {
        "primary": "#00a8ff",
        "primary_dark": "#0097e6",
        "bg_gradient": "linear-gradient(135deg, #000428 0%, #004e92 100%)",
        "secondary_bg": "rgba(255, 255, 255, 0.08)",
        "text": "#f5f6fa"
    },
    "Emerald NeoGlass": {
        "primary": "#00b894",
        "primary_dark": "#00cec9",
        "bg_gradient": "linear-gradient(135deg, #11998e 0%, #38ef7d 100%)",
        "secondary_bg": "rgba(0, 0, 0, 0.2)",
        "text": "#ffffff" # Fixed contrast
    },
    "Sunset Gradient": {
        "primary": "#ff7675",
        "primary_dark": "#d63031",
        "bg_gradient": "linear-gradient(135deg, #ff9a9e 0%, #fecfef 99%, #fecfef 100%)",
        "secondary_bg": "rgba(255,255,255,0.4)",
        "text": "#2d3436"
    },
    "Cyber Neon": {
        "primary": "#0984e3",
        "primary_dark": "#74b9ff",
        "bg_gradient": "radial-gradient(circle at 50% -20%, #3742fa, #000000 90%)",
        "secondary_bg": "rgba(20, 20, 20, 0.8)",
        "text": "#00d2d3"
    },
}

def theme_css():
    theme_name = st.session_state.get("theme", "Midnight Purple")
    cfg = THEME_CONFIG.get(theme_name, THEME_CONFIG["Midnight Purple"])
    
    return f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600&display=swap');
    
    html, body, [class*="css"] {{
        font-family: 'Outfit', sans-serif;
    }}
    
    /* Global App Style */
    .stApp {{
        background: {cfg['bg_gradient']};
        background-attachment: fixed;
        color: {cfg['text']};
    }}
    
    /* Sidebar Glassmorphism */
    section[data-testid="stSidebar"] {{
        background-color: rgba(20, 20, 30, 0.4);
        backdrop-filter: blur(15px);
        -webkit-backdrop-filter: blur(15px);
        border-right: 1px solid rgba(255, 255, 255, 0.05);
    }}
    
    /* Buttons */
    div.stButton > button {{
        background: linear-gradient(90deg, {cfg['primary']}, {cfg['primary_dark']});
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.55rem 1.1rem;
        font-weight: 500;
        transition: all 0.3s cubic-bezier(0.25, 0.8, 0.25, 1);
        box-shadow: 0 4px 10px rgba(0,0,0,0.2);
    }}
    div.stButton > button:hover {{
        transform: translateY(-2px);
        box-shadow: 0 6px 15px rgba(0,0,0,0.3);
    }}
    div.stButton > button:active {{
        transform: translateY(0);
    }}

    /* Inputs */
    div[data-testid="stTextInput"] > div > div {{
        background-color: rgba(255, 255, 255, 0.05);
        color: {cfg['text']};
        border: 1px solid rgba(255, 255, 255, 0.1);
        border-radius: 8px;
    }}
    div[data-testid="stTextInput"] > div > div:focus-within {{
        border-color: {cfg['primary']};
        box-shadow: 0 0 0 1px {cfg['primary']};
    }}
    
    /* Essential Chat Layout */
    .chat-row {{
        display: flex;
        margin: 14px 0;
        gap: 12px;
        width: 100%;
        align-items: flex-start;
    }}
    .chat-row.user {{ flex-direction: row-reverse; }}
    .chat-row.assistant {{ flex-direction: row; }}
    
    /* Bubble visual style */
    .bubble {{
        padding: 12px 18px;
        border-radius: 14px;
        max-width: 80%;
        line-height: 1.5;
        font-size: 16px;
        position: relative;
        box-shadow: 0 2px 5px rgba(0,0,0,0.05);
        animation: fadeIn 0.4s ease-out both;
    }}
    .bubble.user {{
        background: linear-gradient(135deg, {cfg['primary']}, {cfg['primary_dark']});
        color: white;
        border-bottom-right-radius: 2px;
    }}
    .bubble.user p, .bubble.user li, .bubble.user h3, .bubble.user strong, .bubble.user span {{ color: white !important; }}
    
    .bubble p {{ margin: 0 0 6px 0; }}
    .bubble ul, .bubble ol {{ margin: 0 0 6px 0; padding-left: 20px; }}
    .bubble li {{ margin-bottom: 2px; }}
    .bubble strong {{ font-weight: 600; }}
    .bubble h1, .bubble h2, .bubble h3 {{ margin: 8px 0 4px 0; font-size: 1.1em; }}
    .bubble.assistant {{
        background-color: {cfg['secondary_bg']};
        color: {cfg['text']};
        border: 1px solid rgba(255,255,255,0.08);
        border-top-left-radius: 2px;
        backdrop-filter: blur(5px);
    }}
    
    /* Timestamp */
    .timestamp {{
        font-size: 0.7rem;
        opacity: 0.7;
        margin-top: 5px;
        text-align: right;
    }}
    
    /* Animations */
    @keyframes fadeIn {{
        from {{ opacity: 0; transform: translateY(8px); }}
        to {{ opacity: 1; transform: translateY(0); }}
    }}
    
    .typing {{
      display: inline-flex;
      align-items: center;
      gap: 3px;
    }}
    .dot {{
      width: 6px;
      height: 6px;
      background-color: {cfg['text']};
      border-radius: 50%;
      animation: bounce 1.4s infinite ease-in-out both;
      opacity: 0.7;
    }}
    .dot:nth-child(1) {{ animation-delay: -0.32s; }}
    .dot:nth-child(2) {{ animation-delay: -0.16s; }}
    
    @keyframes bounce {{
      0%, 80%, 100% {{ transform: scale(0); }}
      40% {{ transform: scale(1); }}
    }}
    
    /* Headers & Text */
    h1, h2, h3 {{
        font-weight: 600 !important;
        letter-spacing: -0.5px;
        color: {cfg['text']} !important;
    }}
    p, label, span {{
        color: {cfg['text']} !important;
    }}
    </style>
    """

def inject_theme():
    # Always inject the minimal CSS needed for chat layout
    st.markdown(theme_css(), unsafe_allow_html=True)

# -------------------------
# Session defaults
# -------------------------
if "theme" not in st.session_state:
    st.session_state["theme"] = "Midnight Purple"
if "animations" not in st.session_state:
    st.session_state["animations"] = True
if "route" not in st.session_state:
    st.session_state["route"] = "login" if not st.session_state.get("user") else "chat"
if "user" not in st.session_state:
    st.session_state["user"] = None
if "active_cid" not in st.session_state:
    st.session_state["active_cid"] = None
if "chat_mode" not in st.session_state:
    st.session_state["chat_mode"] = "Normal"
if "reset_codes" not in st.session_state:
    st.session_state["reset_codes"] = {}

# Inject theme at start
inject_theme()

# -------------------------
# Diagnostics helper
# -------------------------
def diagnostics_panel():
    try:
        from rag.retriever import get_or_create_index
    except Exception:
        get_or_create_index = None
    with st.sidebar.expander("🛠️ Diagnostics", expanded=False):
        import sys
        st.write(f"Python: `{sys.executable}`")
        files = [f for f in os.listdir(DOCS_DIR) if os.path.splitext(f)[1].lower() in [".pdf", ".txt"]]
        st.write(f"Docs: **{len(files)}**")
        if get_or_create_index:
            try:
                idx = get_or_create_index()
                mat = getattr(idx, "matrix", None)
                chunks = getattr(idx, "chunks", [])
                ready = (mat is not None) and (getattr(mat, "shape", (0,0))[0] > 0) and (len(chunks) > 0)
                st.write(f"Index ready: **{ready}**")
                if ready: st.write(f"Chunks: **{len(chunks)}**")
            except Exception as e:
                st.write("Index: error")
                st.write(str(e))
        else:
            st.write("Retriever not available in this environment")

# -------------------------
# Sidebar: settings, uploads, convos
# -------------------------
def sidebar_ui():
    with st.sidebar:
        st.subheader("⚖️ Legal AI")
        if st.session_state.get("user"):
            u = st.session_state["user"]; email = u["email"]
            st.caption(f"Signed in as **{u['name']}**")
            
            # Theme Selector
            st.markdown("### 🎨 Appearance")
            selected_theme = st.selectbox(
                "Theme", 
                options=list(THEME_CONFIG.keys()), 
                index=list(THEME_CONFIG.keys()).index(st.session_state.get("theme", "Midnight Purple")),
                label_visibility="collapsed"
            )
            if selected_theme != st.session_state.get("theme"):
                st.session_state["theme"] = selected_theme
                st.rerun()

            # Chat Mode Selector
            st.write("")
            st.markdown("### 🧠 Interaction Mode")
            modes = ["Normal", "Summary", "Quiz", "ELI5", "Drafting"]
            selected_mode = st.selectbox(
                "Mode",
                options=modes,
                index=modes.index(st.session_state.get("chat_mode", "Normal")),
                label_visibility="collapsed"
            )
            if selected_mode != st.session_state.get("chat_mode"):
                st.session_state["chat_mode"] = selected_mode
                st.rerun()

            st.divider()
            
            st.markdown("### 💬 Conversations")
            if st.button("➕ New Chat", use_container_width=True, type="primary"):
                st.session_state["active_cid"] = convo_create(email); st.rerun()

            convs = convos_load(email); active = st.session_state.get("active_cid")
            if not convs:
                st.info("No conversations yet.")
            else:
                # render each convo
                for c in convs:
                    # Simple single-column button for selection
                    label = ("🟢 " if c["id"] == active else "• ") + c["title"]
                    if st.sidebar.button(label, key=f"sel_{c['id']}", use_container_width=True):
                        st.session_state["active_cid"] = c["id"]; st.rerun()

            st.divider()
            st.markdown("### 🧰 Chat Tools")
            
            # Tools menu
            c_tools_1, c_tools_2 = st.columns(2)
            with c_tools_1:
                if st.button("🧹 Clear", use_container_width=True, help="Clear chat history"):
                    cid = st.session_state.get("active_cid")
                    if cid: convo_clear(email, cid); st.rerun()
            
            with c_tools_2:
                if st.button("🗑️", use_container_width=True, help="Delete chat"):
                    cid = st.session_state.get("active_cid")
                    if cid:
                         convo_delete(email, cid)
                         remaining = convos_load(email)
                         st.session_state["active_cid"] = remaining[0]["id"] if remaining else None
                         st.rerun()

            # Export functionality
            act = next((c for c in convos_load(email) if c["id"] == st.session_state.get("active_cid")), None)
            if act:
                try:
                    from docx import Document
                    from io import BytesIO
                except ImportError as e:
                    Document = None
                    import_error_msg = str(e)

                if Document:
                    doc = Document()
                    doc.add_heading('Legal Chat Export', 0)
                    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                    doc.add_paragraph(f"Topic: {act.get('title', 'Untitled')}")
                    doc.add_paragraph("-" * 40)
                    
                    for m in act.get("messages", []):
                        role = m.get("role", "unknown").upper()
                        content = m.get("content", "")
                        timestamp = fmt_time(m.get("ts", 0))
                        
                        p = doc.add_paragraph()
                        p.add_run(f"[{timestamp}] {role}:").bold = True
                        doc.add_paragraph(content)
                        doc.add_paragraph("") # Space between messages
                    
                    # Save to buffer
                    bio = BytesIO()
                    doc.save(bio)
                    
                    st.download_button(
                        label="Export Chat (.docx)",
                        data=bio.getvalue(),
                        file_name=f"legal_chat_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                else:
                    st.warning(f"Missing dependency. Error: {import_error_msg}")
                    if st.button("Fix: Install python-docx"):
                        try:
                            import subprocess
                            import sys
                            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
                            st.success("Installed! Please refresh the page (F5).")
                            time.sleep(2)
                            st.rerun()
                        except Exception as install_err:
                            st.error(f"Installation failed: {install_err}")

            st.divider()
            if st.button("Logout", use_container_width=True):
                logout()

        else:
            st.caption("Please sign in to continue.")

        if st.session_state.get("user"):
            st.divider()
            with st.expander("� Manage Documents"):
                files = st.file_uploader("Upload PDF/TXT", type=["pdf","txt"], accept_multiple_files=True)
                if files:
                    saved = 0
                    for f in files:
                        ext = os.path.splitext(f.name)[1].lower()
                        if ext not in [".pdf",".txt"]: continue
                        with open(os.path.join(DOCS_DIR, f.name), "wb") as out:
                            out.write(f.getbuffer()); saved += 1
                    if saved:
                        st.success(f"Uploaded {saved} file(s).")
                        try:
                            from rag.retriever import clear_index_cache
                            clear_index_cache()
                        except Exception:
                            pass
                        st.toast("Index rebuilding...", icon="🧠")
                
                st.write("")
                if st.button("🗑️ Reset Index", use_container_width=True, help="Clear all documents and reset index"):
                    # Clear docs dir
                    for f in os.listdir(DOCS_DIR):
                        fp = os.path.join(DOCS_DIR, f)
                        if os.path.isfile(fp): os.remove(fp)
                    # Clear index
                    try:
                        from rag.retriever import clear_index_cache
                        clear_index_cache()
                    except: pass
                    st.success("Index reset.")
                    time.sleep(1)
                    st.rerun()

            # diagnostics_panel() removed as requested

# -------------------------
# Auth views (centered)
# -------------------------
def make_reset_code(n=6):
    import secrets, string as s
    return "".join(secrets.choice(s.ascii_uppercase + s.digits) for _ in range(n))

def login_view():
    st.write("")
    st.write("")
    
    # Stable Split Layout
    c1, c2 = st.columns([1, 1])
    
    
    with c1:
        # Dynamic Hero Section
        st.markdown("""
        <style>
        .hero-container {
            height: 400px;
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            background: linear-gradient(135deg, rgba(255,255,255,0.1), rgba(255,255,255,0));
            border-radius: 20px;
            padding: 40px;
            border: 1px solid rgba(255,255,255,0.1);
            backdrop-filter: blur(10px);
            text-align: center;
        }
        .hero-icon {
            font-size: 80px;
            margin-bottom: 20px;
            animation: float 6s ease-in-out infinite;
        }
        .hero-title {
            background: linear-gradient(to right, #6c5ce7, #a29bfe);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-weight: 800;
            font-size: 2.5rem;
            margin: 0;
        }
        .hero-subtitle {
            font-size: 1.1rem;
            opacity: 0.8;
            margin-top: 10px;
        }
        @keyframes float {
            0% { transform: translateY(0px); }
            50% { transform: translateY(-20px); }
            100% { transform: translateY(0px); }
        }
        </style>
        <div class="hero-container">
            <div class="hero-icon">⚖️</div>
            <div class="hero-title">Legal AI</div>
            <div class="hero-subtitle">Advanced Retrieval & Analysis</div>
        </div>
        """, unsafe_allow_html=True)
            
    with c2:
        st.markdown("<h1 style='text-align: left; margin-bottom: 5px;'>Welcome Back</h1>", unsafe_allow_html=True)
        st.caption("Sign in to your intelligent legal workspace.")
        st.write("")
        
        email = st.text_input("Email", placeholder="demo@legal.ai")
        pw    = st.text_input("Password", type="password")
        st.write("")
        
        if st.button("Sign In", use_container_width=True, type="primary"):
            if login(email, pw): 
                st.success("Success")
                st.session_state["route"]="chat"
                st.rerun()
            else: 
                st.error("Invalid credentials")

        st.write("") # Minimal unified spacer
        
        c_act1, c_act2 = st.columns(2)
        with c_act1:
            if st.button("Forgot Password?", use_container_width=True): 
                 st.session_state["route"]="forgot"
                 st.rerun()
        with c_act2:
            if st.button("Create Account", use_container_width=True): 
                 st.session_state["route"]="signup"
                 st.rerun()

def signup_view():
    st.write("")
    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        st.markdown("## Create your account")
        st.caption("Start asking legal questions with citations.")
        with st.form("signup_form"):
            name  = st.text_input("Full name")
            email = st.text_input("Email", placeholder="you@example.com")
            pw    = st.text_input("Password", type="password")
            pw2   = st.text_input("Confirm password", type="password")
            create = st.form_submit_button("Create account", use_container_width=True, type="primary")
            if create:
                if not name.strip(): st.error("Please enter your name.")
                elif not email.strip(): st.error("Please enter your email.")
                elif find_user(email): st.error("An account with this email already exists.")
                elif len(pw) < 6: st.error("Password must be at least 6 characters.")
                elif pw != pw2: st.error("Passwords do not match.")
                else:
                    upsert_user({"email": email.strip().lower(), "name": name.strip(), "password_hash": sha256(pw)})
                    st.success("Account created! Signing you in…")
                    login(email, pw); st.session_state["route"]="chat"; st.rerun()
        if st.button("← Back to sign in"): st.session_state["route"]="login"; st.rerun()

def forgot_view():
    st.write("")
    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        st.markdown("## Reset your password")
        st.caption("Enter your email; we’ll generate a reset code (shown here for demo).")
        with st.form("forgot_form"):
            email = st.text_input("Email")
            sent  = st.form_submit_button("Send reset code", use_container_width=True, type="primary")
            if sent:
                if not find_user(email): st.error("No account found for that email.")
                else:
                    code = make_reset_code(); st.session_state.reset_codes[email.lower()] = code
                    st.success("Reset code generated.")
                    st.success("Reset code generated.")
                    st.session_state["pending_reset_email"]=email; st.session_state["route"]="reset"; st.rerun()
        if st.button("← Back to sign in"): st.session_state["route"]="login"; st.rerun()

def reset_view():
    email = st.session_state.get("pending_reset_email")
    if not email: st.session_state["route"]="forgot"; st.rerun()
    st.write("")
    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        st.markdown("## Enter reset code")
        st.caption(f"Email: **{email}**")
        
        # Show code validation for demo
        saved_code = st.session_state.reset_codes.get(email.lower())
        if saved_code:
            st.info(f"ℹ️ **Demo Code:** `{saved_code}`")
            
        with st.form("reset_form"):
            code = st.text_input("Reset code")
            newp = st.text_input("New password", type="password")
            conf = st.text_input("Confirm password", type="password")
            changed = st.form_submit_button("Update password", use_container_width=True, type="primary")
            if changed:
                saved = st.session_state.reset_codes.get(email.lower())
                if not saved or saved.strip()!=code.strip(): st.error("Invalid or expired code.")
                elif len(newp)<6: st.error("Password must be at least 6 characters.")
                elif newp!=conf: st.error("Passwords do not match.")
                else:
                    u = find_user(email)
                    if not u: st.error("Account not found.")
                    else:
                        u["password_hash"]=sha256(newp); upsert_user(u)
                        st.session_state.reset_codes.pop(email.lower(), None)
                        st.success("Password updated! Please sign in.")
                        st.session_state.pop("pending_reset_email", None)
                        st.session_state["route"]="login"; st.rerun()
        if st.button("← Back"): st.session_state["route"]="forgot"; st.rerun()

# -------------------------
# Chat helpers & rendering
# -------------------------
def render_bubble(role: str, content: str, ts: float = None, animate: bool = True, suggestions: list = None, key_prefix: str = ""):
    # Use markdown to render content to HTML
    try:
        content_html = markdown.markdown(content)
    except Exception:
        content_html = esc(content).replace("\n", "<br>")
    
    timestamp = f'<div class="timestamp">{fmt_time(ts)}</div>' if ts else ""
    if role == "user":
        cls = "bubble user"
        if animate and st.session_state.get("animations", True):
            cls += " anim-right send-ripple"
        st.markdown(f'<div class="chat-row user"><div class="{cls}">{content_html}{timestamp}</div></div>', unsafe_allow_html=True)
    else:
        cls = "bubble assistant"
        if animate and st.session_state.get("animations", True):
            cls += " anim-left glow"
        st.markdown(f'<div class="chat-row assistant"><div class="{cls}">{content_html}{timestamp}</div></div>', unsafe_allow_html=True)
        
        # Render specific follow-up suggestions if present
        if suggestions:
            st.write("")
            display_suggestions(suggestions, key_prefix)

def render_animated_background():
    # Subtle floating orbs in the background
    st.markdown("""
    <style>
    .stApp {
        background: transparent !important; /* Ensure default Streamlit bg doesn't block */
    }
    .animated-bg-container {
        position: fixed;
        top: 0;
        left: 0;
        width: 100vw;
        height: 100vh;
        z-index: -1;
        overflow: hidden;
        background: linear-gradient(135deg, #0f0c29 0%, #302b63 50%, #24243e 100%); /* Base */
    }
    .orb {
        position: absolute;
        border-radius: 50%;
        filter: blur(60px);
        opacity: 0.5;
        animation: floatNav 15s infinite ease-in-out alternate;
    }
    .orb-1 {
        top: -10%;
        left: -10%;
        width: 600px;
        height: 600px;
        background: radial-gradient(circle, rgba(108, 92, 231, 0.4), transparent);
    }
    .orb-2 {
        bottom: -10%;
        right: -10%;
        width: 500px;
        height: 500px;
        background: radial-gradient(circle, rgba(0, 210, 211, 0.3), transparent);
        animation-delay: -5s;
    }
    .orb-3 {
        top: 40%;
        left: 40%;
        width: 300px;
        height: 300px;
        background: radial-gradient(circle, rgba(255, 118, 117, 0.2), transparent);
        animation-duration: 25s;
    }
    @keyframes floatNav {
        0% { transform: translate(0, 0) scale(1); }
        100% { transform: translate(30px, 40px) scale(1.1); }
    }
    
    /* Animated Header Text */
    .radiant-text {
        background: linear-gradient(to right, #fff, #a29bfe, #74b9ff, #fff);
        background-size: 200% auto;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        animation: shine 5s linear infinite;
        font-weight: 800;
        font-size: 2.5rem;
        margin-bottom: 0.5rem;
    }
    @keyframes shine {
        to { background-position: 200% center; }
    }
    </style>
    <div class="animated-bg-container">
        <div class="orb orb-1"></div>
        <div class="orb orb-2"></div>
        <div class="orb orb-3"></div>
    </div>
    """, unsafe_allow_html=True)

def display_suggestions(suggestions, key_prefix):
    # Use small columns for pills
    cols = st.columns(len(suggestions))
    for i, sug in enumerate(suggestions):
        # Add visual numbering
        label = f"{i+1}. {sug}"
        if cols[i].button(label, key=f"sug_{key_prefix}_{i}", use_container_width=True):
            st.session_state["suggestion_click"] = sug
            st.rerun()

def chat_view():
    render_animated_background()
    
    # Custom animated header
    st.markdown('<div class="radiant-text">AI Legal Assistant</div>', unsafe_allow_html=True)
    st.caption("Advanced Retrieval & Reference System")

    user = st.session_state.get("user")
    if not user:
        st.error("Please sign in to use the assistant.")
        return

    email = user["email"]
    convs = convos_load(email)
    if not convs:
        st.session_state["active_cid"] = convo_create(email)
        convs = convos_load(email)

    active_cid = st.session_state.get("active_cid") or (convs[0]["id"] if convs else convo_create(email))
    active = next((c for c in convs if c["id"] == active_cid), None)
    if active is None:
        active = convs[0] if convs else None
        st.session_state["active_cid"] = active["id"] if active else None

    if active:
        # Check for suggestion click from previous rerun
        if "suggestion_click" in st.session_state:
            sc = st.session_state.pop("suggestion_click")
            convo_append_msg(email, active["id"], "user", sc)
            st.rerun()

        # Render messages
        for i, m in enumerate(active["messages"]):
            is_last = (i == len(active["messages"]) - 1)
            sugs = m.get("suggestions") if is_last else None
            render_bubble(m["role"], m["content"], ts=m.get("ts"), animate=False, suggestions=sugs, key_prefix=f"{active['id']}_{i}")

        # If conversation is new (only 1 welcome message), show animated starter cards
        if len(active["messages"]) <= 1:
            st.markdown("""
            <style>
            .starter-card-container {
                animation: slideUpCard 0.8s cubic-bezier(0.2, 0.8, 0.2, 1);
            }
            @keyframes slideUpCard {
                from { opacity: 0; transform: translateY(30px); }
                to { opacity: 1; transform: translateY(0); }
            }
            </style>
            """, unsafe_allow_html=True)
            
            st.write("")
            st.write("")
            st.markdown('<div class="starter-card-container"></div>', unsafe_allow_html=True) # Animation trigger
            
            c1, c2, c3 = st.columns(3)
            starters = [
                ("📜 Legal Rights", "What are my rights if arrested?"),
                ("🏘️ Property Law", "Draft a rental agreement clause"),
                ("🚗 Traffic Rules", "Fine for driving without license?"),
                ("💼 Corporate", "Process to incorporate a startup?"),
                ("🧠 IP Rights", "How to trademark a logo?"),
                ("⚖️ Criminal", "What is the punishment for theft?")
            ]
            
            for i, (emoji, q_text) in enumerate(starters):
                col = [c1, c2, c3][i % 3]
                if col.button(f"{emoji}\n\n{q_text}", key=f"start_{i}", use_container_width=True):
                    convo_append_msg(email, active["id"], "user", q_text)
                    st.rerun()

    # Input handling
    prompt = st.chat_input("Type your legal question…")
    if prompt and active:
        # Numeric shortcut handling
        if prompt.strip().isdigit():
            idx = int(prompt.strip()) - 1
            # Check if last message was assistant and has suggestions
            if active["messages"] and active["messages"][-1]["role"] == "assistant":
                last_sugs = active["messages"][-1].get("suggestions", [])
                if 0 <= idx < len(last_sugs):
                    prompt = last_sugs[idx] # Swap number for question text
        
        convo_append_msg(email, active["id"], "user", prompt)
        st.rerun()

    # Reaction Loop: If last message is User, generate Assistant response
    if active and active["messages"] and active["messages"][-1]["role"] == "user":
        # Show typing indicator
        placeholder = st.empty()
        typing_html = (
            '<div class="chat-row assistant"><div class="bubble assistant anim-left">'
            '<span class="typing"><span class="dot"></span><span class="dot"></span><span class="dot"></span></span>'
            '</div></div>'
        )
        if st.session_state.get("animations", True):
            placeholder.markdown(typing_html, unsafe_allow_html=True)
        else:
            placeholder.markdown('<div class="chat-row assistant"><div class="bubble assistant">Searching…</div></div>', unsafe_allow_html=True)

        # Generate response
        user_query = active["messages"][-1]["content"]
        try:
            mode = st.session_state.get("chat_mode", "Normal")
            resp = answer_query(user_query, mode=mode)
            
            if isinstance(resp, str):
                reply = resp
                sugs = []
            else:
                reply = resp.get("content", "")
                sugs = resp.get("suggestions", [])
                
        except Exception as e:
            reply = f"Error: {e}"
            sugs = []

        placeholder.empty()
        
        # Append and show
        convo_append_msg(email, active["id"], "assistant", reply, suggestions=sugs)
        
        # Force rerun to render the new message naturally in the loop
        st.rerun()

# -------------------------
# Router
# -------------------------
if "route" not in st.session_state:
    st.session_state["route"] = "login" if not st.session_state.get("user") else "chat"

def main():
    sidebar_ui()

    route = st.session_state.get("route","login")
    if not st.session_state.get("user"):
        if route=="login":   login_view()
        elif route=="forgot": forgot_view()
        elif route=="reset":  reset_view()
        elif route=="signup": signup_view()
        else: st.session_state["route"]="login"; st.rerun()
    else:
        if route!="chat": st.session_state["route"]="chat"; st.rerun()
        chat_view()

if __name__ == "__main__":
    main()
